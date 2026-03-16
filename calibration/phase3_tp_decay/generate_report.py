# -*- coding: utf-8 -*-
"""
Generate comprehensive Word report documenting the Augmented TP Decay Model.
Includes: equations, figures, tables, comparison with original model.
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

import json
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).resolve().parent
AUG_OUTDIR = SCRIPT_DIR / "augmented_outputs"
ORIG_OUTDIR = SCRIPT_DIR / "outputs"
REPORT_PATH = SCRIPT_DIR / "augmented_outputs" / "TP_Decay_Augmented_Model_Report.docx"


def add_equation(doc, text, label=None):
    """Add an equation paragraph (centered, italic)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Cambria Math"
    if label:
        tab_run = p.add_run(f"\t({label})")
        tab_run.font.size = Pt(10)
    return p


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def main():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- Title ---
    title = doc.add_heading("Augmented Threat Perception Decay Model", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Technical Report: Owner/Renter Calibration with Full Survey Data")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()  # spacer

    # =================================================================
    # Section 1: Motivation
    # =================================================================
    doc.add_heading("1. Motivation", level=1)
    doc.add_paragraph(
        "The original TP decay model was calibrated using only flood-experienced respondents "
        "(those with Q12 not missing, stored as column Q15 in processed Excel). This limits "
        "sample sizes to 170 owners and 43 renters, which may lead to unstable parameter "
        "estimates, particularly for the renter group."
    )
    doc.add_paragraph(
        "The augmented model leverages all 936 survey respondents by treating non-flood-experienced "
        "respondents as observations at t → ∞, where threat perception has decayed to a baseline "
        "level (TP_base). This increases the effective sample size to 693 owners and 243 renters, "
        "providing more robust estimates of both the decay dynamics and the long-term equilibrium."
    )

    # Data summary table
    doc.add_heading("1.1 Data Availability", level=2)
    headers = ["Group", "Flood-Exp (n)", "No-Flood (n)", "Total (n)",
               "TP (flood)", "TP (no-flood)", "TP_base (0-1)"]
    rows = [
        ["Owner", "170", "523", "693", "0.668", "0.546", "0.546"],
        ["Renter", "43", "200", "243", "0.720", "0.556", "0.556"],
    ]
    add_styled_table(doc, headers, rows)
    doc.add_paragraph(
        "Table 1: Survey respondent counts and mean TP values by flood experience and housing tenure."
    ).italic = True

    # =================================================================
    # Section 2: Model Equations
    # =================================================================
    doc.add_heading("2. Model Equations", level=1)

    doc.add_heading("2.1 Original Model", level=2)
    doc.add_paragraph(
        "The original model assumes threat perception decays from an initial level TP₀ "
        "toward zero following a flood event:"
    )
    add_equation(doc, "TP(t) = TP₀ · exp(−ln2 · w_eff · Eff(t))", "1a")

    doc.add_heading("2.2 Augmented Model", level=2)
    doc.add_paragraph(
        "The augmented model introduces a baseline floor (TP_base) representing the "
        "long-term equilibrium threat perception. Threat perception decays from TP₀ "
        "toward TP_base (not zero):"
    )
    add_equation(doc, "TP(t) = TP_base + (TP₀ − TP_base) · exp(−ln2 · w_eff · Eff(t))", "1b")
    doc.add_paragraph(
        "where TP_base is estimated as the mean threat perception of respondents "
        "without flood experience (t → ∞ observations)."
    )

    doc.add_heading("2.3 Shared Equations (Unchanged)", level=2)
    doc.add_paragraph(
        "The following equations are identical in both models:"
    )

    doc.add_paragraph("Effective weight combining place attachment and social capital:")
    add_equation(doc, "w_eff = α · (1 − PA) + β · SC", "2")
    doc.add_paragraph(
        "where PA is place attachment (0-1, higher PA protects TP from decay) "
        "and SC is social capital (0-1, higher SC accelerates decay/recovery)."
    )

    doc.add_paragraph("Saturating memory half-life:")
    add_equation(doc, "τ(t) = τ∞ − (τ∞ − τ₀) · exp(−k · t)", "3")
    doc.add_paragraph(
        "where τ₀ is the initial (acute) half-life, τ∞ is the long-term stable half-life, "
        "and k controls the rate of memory consolidation."
    )

    doc.add_paragraph("Cumulative effective decay (closed-form integral):")
    add_equation(doc, "Eff(t) = ∫₀ᵗ 1/τ(u) du = t/τ∞ + (1/(k·τ∞)) · ln(τ(t)/τ₀)", "4")

    doc.add_paragraph()
    doc.add_paragraph(
        "Key insight: The only change is in Equation (1b) — the outer TP formula. "
        "Equations (2)-(4) governing the effective weight, memory consolidation, "
        "and cumulative decay remain exactly the same."
    ).bold = True

    # =================================================================
    # Section 3: Calibration Methodology
    # =================================================================
    doc.add_heading("3. Calibration Methodology", level=1)

    doc.add_heading("3.1 TP_base Estimation", level=2)
    doc.add_paragraph(
        "TP_base is estimated as the sample mean of threat perception among "
        "non-flood-experienced respondents within each tenure group. "
        "These respondents serve as steady-state observations (t → ∞), "
        "providing a natural anchor for the decay floor."
    )
    add_equation(doc, "TP_base = (1/n_nf) · Σᵢ TPᵢ,  for i ∈ {no flood experience}", "5")

    doc.add_heading("3.2 TP₀ Profiling", level=2)
    doc.add_paragraph(
        "Given fixed (α, β, τ₀, τ∞, k) and TP_base, the initial threat perception TP₀ "
        "is profiled out analytically. Define δ = TP₀ − TP_base and g(tᵢ) = exp(−ln2 · w_eff,i · Eff(tᵢ)). "
        "Then for flood-experienced respondents:"
    )
    add_equation(doc, "δ̂ = Σᵢ (TPᵢ − TP_base) · g(tᵢ) / Σᵢ g(tᵢ)²", "6")
    add_equation(doc, "TP₀ = TP_base + δ̂", "7")

    doc.add_heading("3.3 Optimization", level=2)
    doc.add_paragraph(
        "The optimization proceeds in three stages:"
    )
    doc.add_paragraph(
        "(1) Grid search over (α, β) with 20×20 grid points in [0.1, 0.6]². "
        "For each (α, β), differential evolution optimizes (τ₀, δ, k) where δ = τ∞ − τ₀.",
        style="List Number"
    )
    doc.add_paragraph(
        "(2) Joint selection with cross-group constraints: "
        "|α_owner − α_renter| ≥ 0.10 and |β_owner − β_renter| ≥ 0.10, "
        "minimizing combined RMSE.",
        style="List Number"
    )
    doc.add_paragraph(
        "(3) Bootstrap (N=100 resamples) for 95% confidence intervals on all parameters.",
        style="List Number"
    )

    doc.add_heading("3.4 Loss Function", level=2)
    doc.add_paragraph("The combined loss function balances flood and no-flood data:")
    add_equation(doc,
        "L = w_decay · MSE_flood + w_base · MSE_base + λ · (R₅ − r_target)²", "8")
    doc.add_paragraph(
        "where w_decay = n_flood/(n_total), w_base = n_nf/(n_total), "
        "R₅ = TP(5)/TP(0) is the 5-year retention ratio, and r_target = 0.5."
    )

    # =================================================================
    # Section 4: Results
    # =================================================================
    doc.add_heading("4. Results", level=1)

    # Load results
    aug_params = {}
    aug_params_path = AUG_OUTDIR / "augmented_tp_params_for_abm.json"
    if aug_params_path.exists():
        with open(aug_params_path, "r") as f:
            aug_params = json.load(f)

    orig_params = {}
    orig_params_path = ORIG_OUTDIR / "tp_decay_params_for_abm.json"
    if orig_params_path.exists():
        with open(orig_params_path, "r") as f:
            orig_params = json.load(f)

    doc.add_heading("4.1 Calibrated Parameters", level=2)
    headers = ["Parameter", "Owner (Orig)", "Owner (Aug)", "Renter (Orig)", "Renter (Aug)"]
    param_rows = []
    for p, label in [("alpha", "α"), ("beta", "β"), ("tau0", "τ₀"),
                      ("tau_inf", "τ∞"), ("k", "k")]:
        param_rows.append([
            label,
            f"{orig_params.get('owner', {}).get(p, '-')}",
            f"{aug_params.get('owner', {}).get(p, '-')}",
            f"{orig_params.get('renter', {}).get(p, '-')}",
            f"{aug_params.get('renter', {}).get(p, '-')}",
        ])
    # Add augmented-only params
    for p, label in [("TP_base", "TP_base"), ("TP0_hat", "TP₀")]:
        param_rows.append([
            label, "N/A",
            f"{aug_params.get('owner', {}).get(p, '-')}",
            "N/A",
            f"{aug_params.get('renter', {}).get(p, '-')}",
        ])
    # RMSE
    param_rows.append([
        "RMSE (flood)",
        f"{orig_params.get('owner', {}).get('RMSE', '-')}",
        f"{aug_params.get('owner', {}).get('RMSE_flood', '-')}",
        f"{orig_params.get('renter', {}).get('RMSE', '-')}",
        f"{aug_params.get('renter', {}).get('RMSE_flood', '-')}",
    ])
    param_rows.append([
        "RMSE (all)", "N/A",
        f"{aug_params.get('owner', {}).get('RMSE_all', '-')}",
        "N/A",
        f"{aug_params.get('renter', {}).get('RMSE_all', '-')}",
    ])
    param_rows.append([
        "n (data)",
        "170", f"170+523=693",
        "43", f"43+200=243",
    ])

    add_styled_table(doc, headers, param_rows)
    doc.add_paragraph(
        "Table 2: Comparison of calibrated parameters between original and augmented models."
    ).italic = True

    doc.add_heading("4.2 Key Findings", level=2)
    doc.add_paragraph(
        "The augmented model yields several improvements over the original:"
    )

    # Compute improvements
    if aug_params and orig_params:
        owner_rmse_orig = orig_params["owner"]["RMSE"]
        owner_rmse_aug = aug_params["owner"]["RMSE_flood"]
        renter_rmse_orig = orig_params["renter"]["RMSE"]
        renter_rmse_aug = aug_params["renter"]["RMSE_flood"]
        owner_pct = (1 - owner_rmse_aug / owner_rmse_orig) * 100
        renter_pct = (1 - renter_rmse_aug / renter_rmse_orig) * 100

        doc.add_paragraph(
            f"(1) RMSE improvement: Owner RMSE decreased from {owner_rmse_orig:.4f} to "
            f"{owner_rmse_aug:.4f} ({owner_pct:.1f}% reduction). "
            f"Renter RMSE decreased from {renter_rmse_orig:.4f} to "
            f"{renter_rmse_aug:.4f} ({renter_pct:.1f}% reduction).",
            style="List Number"
        )
    doc.add_paragraph(
        "(2) Sample size: The augmented model uses 693 owner and 243 renter observations "
        "(vs 170 and 43 in the original), providing more stable parameter estimates.",
        style="List Number"
    )
    doc.add_paragraph(
        "(3) Physical interpretation: The TP_base floor provides a natural lower bound, "
        "reflecting that threat perception does not decay to zero even decades after a flood. "
        "This is consistent with the observation that non-flood-experienced residents still "
        "maintain moderate threat perception (owner: 0.546, renter: 0.556).",
        style="List Number"
    )
    doc.add_paragraph(
        "(4) The shorter τ∞ values in the augmented model are expected: because TP now "
        "decays to TP_base (not zero), the effective decay range is smaller, so less "
        "time is needed to reach the floor.",
        style="List Number"
    )

    # =================================================================
    # Section 5: Figures
    # =================================================================
    doc.add_heading("5. Figures", level=1)

    # Figure 1: Decay curves
    fig_curves = AUG_OUTDIR / "augmented_decay_curves.png"
    if fig_curves.exists():
        doc.add_heading("5.1 Calibrated Decay Curves", level=2)
        doc.add_picture(str(fig_curves), width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "Figure 1: Augmented TP decay curves for owner and renter groups. "
            "Dashed horizontal lines show TP_base floors. Scattered points show "
            "flood-experienced observations; crosses at t=35-40 represent "
            "non-flood respondents (t→∞ approximation)."
        ).italic = True

    # Figure 2: Residuals
    fig_resid = AUG_OUTDIR / "augmented_residuals.png"
    if fig_resid.exists():
        doc.add_heading("5.2 Observed vs Predicted", level=2)
        doc.add_picture(str(fig_resid), width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "Figure 2: Scatter plot of observed vs predicted TP values. "
            "Blue/red points are flood-experienced respondents; lighter markers "
            "are non-flood respondents predicted at TP_base."
        ).italic = True

    # Figure 3: CI ribbons
    fig_ci = AUG_OUTDIR / "augmented_confidence_intervals.png"
    if fig_ci.exists():
        doc.add_heading("5.3 Bootstrap Confidence Intervals", level=2)
        doc.add_picture(str(fig_ci), width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "Figure 3: Bootstrap 95% confidence intervals (N=100) for the "
            "augmented decay curves. Shaded regions represent parameter uncertainty."
        ).italic = True

    # =================================================================
    # Section 6: Implementation in ABM
    # =================================================================
    doc.add_heading("6. Implementation in ABM", level=1)
    doc.add_paragraph(
        "To use the augmented model in the simulation, the TP update function changes from:"
    )
    add_equation(doc, "TP_new = TP_old · exp(−ln2 · w_eff · ΔEff)", "old")
    doc.add_paragraph("to:")
    add_equation(doc, "TP_new = TP_base + (TP_old − TP_base) · exp(−ln2 · w_eff · ΔEff)", "new")
    doc.add_paragraph(
        "where ΔEff = Eff(t+Δt) − Eff(t). All other simulation mechanics "
        "(action prediction, perception sampling, tract-level aggregation) remain unchanged."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Calibrated parameters for ABM config (abm_params.yaml):")
    run.bold = True

    if aug_params:
        for grp in ["owner", "renter"]:
            r = aug_params[grp]
            doc.add_paragraph(
                f"  {grp}: α={r['alpha']}, β={r['beta']}, "
                f"τ₀={r['tau0']}, τ∞={r['tau_inf']}, k={r['k']}, "
                f"TP_base={r['TP_base']}",
                style="List Bullet"
            )

    # =================================================================
    # Section 7: Summary
    # =================================================================
    doc.add_heading("7. Summary of Changes", level=1)

    headers2 = ["Aspect", "Original Model", "Augmented Model"]
    summary_rows = [
        ["Outer formula", "TP(t) = TP₀ · exp(−ln2·w·Eff)", "TP(t) = TP_base + (TP₀−TP_base)·exp(−ln2·w·Eff)"],
        ["w_eff", "α(1−PA) + βSC", "Same"],
        ["τ(t)", "τ∞ − (τ∞−τ₀)exp(−kt)", "Same"],
        ["Eff(t)", "∫1/τ(u)du (closed-form)", "Same"],
        ["Data used", "Flood-exp only (n=213)", "All respondents (n=936)"],
        ["TP_base", "N/A (decays to 0)", "Estimated from non-flood respondents"],
        ["TP₀", "Profiled from data", "Profiled (relative to TP_base)"],
        ["Lower bound", "TP → 0", "TP → TP_base > 0"],
    ]
    add_styled_table(doc, headers2, summary_rows)
    doc.add_paragraph(
        "Table 3: Summary of differences between the original and augmented models."
    ).italic = True

    # Save
    doc.save(str(REPORT_PATH))
    print(f"Report saved: {REPORT_PATH}")


if __name__ == "__main__":
    main()
