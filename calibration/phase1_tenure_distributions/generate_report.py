"""
Generate Phase 1 Report: Owner/Renter Data Extraction & Distribution Fitting
Output: Word document with tables, embedded plots, and analysis.
Data source: 999_value.xlsx (raw Qualtrics, n=936 after QC)
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

OUT_DIR = Path(r"C:\Users\wenyu\OneDrive - Lehigh University\Desktop\Lehigh\NSF-project\ABM\Calibration\tenure_distributions")

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return h

def add_table_from_df(df, col_widths=None):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, val in enumerate(row):
            row_cells[j].text = str(val)
            for p in row_cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# ══════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════
title = doc.add_heading("Phase 1 Report: Owner/Renter Data Extraction\n& Distribution Fitting", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph("Date: 2026-02-26 | Data source: 999_value.xlsx (raw Qualtrics export)")
doc.add_paragraph("")

# ══════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════
add_heading("1. Overview", level=1)
doc.add_paragraph(
    "This report documents the extraction and distribution fitting of survey data "
    "for the ABM's homeowner vs renter (housing tenure) grouping axis. "
    "Data is drawn from the original Qualtrics export (999_value.xlsx, 1039 respondents). "
    "Classification rule: Q5 = 1 → renter; Q5 ∈ {2, 3} → owner."
)

add_heading("1.1 Data Source & QC Filters", level=2)
doc.add_paragraph(
    "• Raw data: 999_value.xlsx (1039 respondents, 108 columns)\n"
    "• Filter 1: Finished = 1 (completed survey)\n"
    "• Filter 2: Q5 ∈ {1, 2, 3} (valid housing status)\n"
    "• Filter 3: Q23 = 3 (attention check pass — 'Neutral')\n"
    "• Filter 4: Drop 30 lowest-quality owners (straightlining + fast completion)\n"
    "• Q5 encoding: 1 = renter, 2 = owner (with mortgage), 3 = owner (without mortgage)"
)

add_heading("1.2 Variable Formulas (verified from raw survey)", level=2)
formula_data = pd.DataFrame({
    "Variable": ["TP_mean", "CP_mean", "SP_mean", "SC_mean", "PA_mean",
                  "FI", "EH", "BP", "RL"],
    "Formula": [
        "mean(Q22_1 : Q22_11) — 11 items",
        "mean(Q24_1, Q24_2, Q25_1, Q25_2, Q25_4, Q25_5, Q25_7, Q25_8) — 8 items",
        "mean(Q25_3, Q25_6, Q25_9) — 3 items",
        "mean(Q21_1 : Q21_6) — 6 items",
        "mean(Q21_7 : Q21_15) — 9 items",
        "Q27", "Q29 (owner only)", "Q31 (owner only)", "Q33 (renter only)"
    ],
    "Original Q#": [
        "Q19_1-Q19_11", "Q21_1-2 + Q22_1,2,4,5,7,8", "Q22_3, Q22_6, Q22_9",
        "Q18_1-Q18_6", "Q18_7-Q18_15",
        "Q24", "Q26", "Q28", "Q30"
    ],
    "Scale": ["1–5 Likert", "1–5 Likert", "1–5 Likert", "1–5 Likert", "1–5 Likert",
              "1–5 Likert", "1–5 Likert", "1–5 Likert", "1–5 Likert"]
})
add_table_from_df(formula_data)
doc.add_paragraph("")

# ══════════════════════════════════════════════════════════════════
# 2. SAMPLE SIZES
# ══════════════════════════════════════════════════════════════════
add_heading("2. Sample Sizes After Filtering", level=1)

size_data = pd.DataFrame({
    "Dataset": ["Full Survey", "Full Survey", "Full Survey",
                "TP Decay Calibration", "TP Decay Calibration", "TP Decay Calibration"],
    "Group": ["Owner", "Renter", "Total", "Owner", "Renter", "Total"],
    "n": [557, 379, 936, 135, 73, 208],
})
add_table_from_df(size_data)
doc.add_paragraph("")
doc.add_paragraph(
    "The TP decay calibration subset includes only respondents with flood experience "
    "(Q15 not NaN). The renter calibration sample (n=73) is adequate for grid-search "
    "calibration and substantially larger than the previous owner/renter split (n=43)."
)

# ══════════════════════════════════════════════════════════════════
# 3. BETA DISTRIBUTION FITS
# ══════════════════════════════════════════════════════════════════
add_heading("3. Beta Distribution Parameters", level=1)
doc.add_paragraph(
    "Each psychological variable was normalized to [0, 1] by dividing by 5 (Likert 1–5 → 0.2–1.0), "
    "then fitted with a Beta(α, β) distribution using MLE (scipy.stats.beta.fit with floc=0, fscale=1). "
    "The KS test checks goodness-of-fit (p > 0.05 = good fit)."
)

beta_df = pd.read_csv(OUT_DIR / "beta_parameters_summary.csv")
add_table_from_df(beta_df)
doc.add_paragraph("")

add_heading("3.1 Key Observations", level=2)
doc.add_paragraph(
    "• TP (Threat Perception): Both groups show similar means (~0.58). "
    "Owner has tighter distribution (α=6.27, β=4.64). "
    "Both pass KS test (p>0.10).\n"
    "• CP (Coping Perception): Owner has higher concentration (α=7.86) vs renter (α=3.82). "
    "Means are similar (~0.61–0.64).\n"
    "• SP (Social Perception): Renter has wider spread (α=1.76) vs owner (α=3.40). "
    "3-item scale produces sparse discrete distribution.\n"
    "• SC (Social Capital): Both groups strongly right-skewed (high SC). "
    "Owner mean 0.74, renter mean 0.73. 6-item scale.\n"
    "• PA (Place Attachment): Owner has tighter distribution (α=5.54) vs renter (α=2.51). "
    "Similar means (~0.65).\n"
    "• KS test: TP passes for both groups. Other variables show significant deviation — "
    "expected for discrete Likert data. Beta approximation is standard for ABM initialization."
)

# ══════════════════════════════════════════════════════════════════
# 4. ACTION ADOPTION RATES
# ══════════════════════════════════════════════════════════════════
add_heading("4. Action Adoption Rates", level=1)
action_df = pd.read_csv(OUT_DIR / "action_adoption_rates.csv")
add_table_from_df(action_df)
doc.add_paragraph("")

add_heading("4.1 Key Observations", level=2)
doc.add_paragraph(
    "• FI (Flood Insurance, Q24): 100% response rate for both groups. "
    "Renters have slightly higher mean intention (3.01 vs 2.86).\n"
    "• EH (Elevation/Hardening, Q26): Owner-only action, 100% valid responses (n=557). "
    "Mean intention 2.61.\n"
    "• BP (Building Protection, Q28): Owner-only action, 100% valid responses (n=557). "
    "Mean intention 2.39.\n"
    "• RL (Relocation, Q30): Renter-only action, 100% valid responses (n=379). "
    "Mean intention 3.13.\n\n"
    "The tenure-gated skip logic in the survey ensures clean separation: "
    "owners answer EH/BP but not RL, renters answer RL but not EH/BP. "
    "All groups answer FI."
)

# ══════════════════════════════════════════════════════════════════
# 5. PLOTS
# ══════════════════════════════════════════════════════════════════
add_heading("5. Distribution Plots", level=1)

add_heading("5.1 Overlay Comparison (Owner vs Renter)", level=2)
overlay_img = OUT_DIR / "overlay_comparison.png"
if overlay_img.exists():
    doc.add_picture(str(overlay_img), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph("[Plot not generated yet — run extract_tenure_data.py first]")

add_heading("5.2 Beta Fits with Histograms", level=2)
beta_img = OUT_DIR / "beta_distributions_owner_renter.png"
if beta_img.exists():
    doc.add_picture(str(beta_img), width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph("[Plot not generated yet — run extract_tenure_data.py first]")

add_heading("5.3 Action Response Rates", level=2)
action_rate_img = OUT_DIR / "action_rates_owner_renter.png"
if action_rate_img.exists():
    doc.add_picture(str(action_rate_img), width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph("[Plot not generated yet — run extract_tenure_data.py first]")

add_heading("5.4 Action Mean Likert Scores", level=2)
action_mean_img = OUT_DIR / "action_means_owner_renter.png"
if action_mean_img.exists():
    doc.add_picture(str(action_mean_img), width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph("[Plot not generated yet — run extract_tenure_data.py first]")

# ══════════════════════════════════════════════════════════════════
# 6. TP DECAY CALIBRATION DATA SUMMARY
# ══════════════════════════════════════════════════════════════════
add_heading("6. TP Decay Calibration Data", level=1)
doc.add_paragraph(
    "The calibration dataset contains only respondents with flood experience (Q15 not NaN). "
    "Variables include Q15 (flood recency ordinal, original survey Q12), Q15_year "
    "(converted to approximate years), and the 5 psychological means."
)

cal_summary = pd.DataFrame({
    "Statistic": ["Count", "TP mean", "CP mean", "SP mean", "SC mean", "PA mean"],
    "Owner (n=135)": ["135", "—", "—", "—", "—", "—"],
    "Renter (n=73)": ["73", "—", "—", "—", "—", "—"],
})
add_table_from_df(cal_summary)
doc.add_paragraph("")
doc.add_paragraph(
    "Note: Exact descriptive statistics will be computed during Phase 3 calibration. "
    "The renter calibration sample (n=73) is substantially larger than the previous "
    "owner/renter split (n=43), which should improve TP decay parameter estimates."
)

# ══════════════════════════════════════════════════════════════════
# 7. READY-TO-USE PARAMETERS
# ══════════════════════════════════════════════════════════════════
add_heading("7. Ready-to-Use PerceptionSampler Parameters", level=1)
doc.add_paragraph(
    "These Beta parameters are used in tp.py (BETA_PARAMS_OWNER_DEFAULT / "
    "BETA_PARAMS_RENTER_DEFAULT) to initialize ~52,000 household agents:"
)

code = doc.add_paragraph()
code.style = doc.styles["Normal"]
code_text = """params = {
    'owner': {  # n=557
        'TP': {'alpha': 6.270198, 'beta': 4.635320},
        'CP': {'alpha': 7.857397, 'beta': 5.051820},
        'SP': {'alpha': 3.398335, 'beta': 2.752034},
        'SC': {'alpha': 4.773929, 'beta': 1.689776},
        'PA': {'alpha': 5.539009, 'beta': 2.941853},
    },
    'renter': {  # n=379
        'TP': {'alpha': 3.713441, 'beta': 2.668174},
        'CP': {'alpha': 3.824367, 'beta': 2.164083},
        'SP': {'alpha': 1.759103, 'beta': 1.103402},
        'SC': {'alpha': 2.794319, 'beta': 1.035687},
        'PA': {'alpha': 2.512725, 'beta': 1.322463},
    }
}"""
run = code.add_run(code_text)
run.font.name = "Consolas"
run.font.size = Pt(9)

# ══════════════════════════════════════════════════════════════════
# 8. NEXT STEPS
# ══════════════════════════════════════════════════════════════════
add_heading("8. Next Steps", level=1)
doc.add_paragraph(
    "1. Professor reviews this report and approves distributions.\n"
    "2. Phase 2: Re-run Bayesian Beta regression with new data (owner n=557, renter n=379). "
    "Models: owner×{FI,EH,BP} + renter×{FI,RL}.\n"
    "3. Phase 3: Re-run TP decay calibration with new cal data "
    "(owner_cal n=135, renter_cal n=73).\n"
    "4. Update simulation code with new parameters."
)

# ── Save ──
out_path = OUT_DIR / "Phase1_Owner_Renter_Report.docx"
doc.save(str(out_path))
print(f"Report saved to: {out_path}")
